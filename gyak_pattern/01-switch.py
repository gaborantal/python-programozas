from docx import Document
import os
from striprtf.striprtf import rtf_to_text


def read_and_print_3(melyik_file):
    file_path = os.path.join("resources", melyik_file)

    match melyik_file.split("."):
        case [_, "txt"]:
            with open(file_path, encoding="utf8") as fp:
                print(fp.read())
        case [_, "docx"]:
            doc = Document(file_path)
            for para in doc.paragraphs:
                print(para.text)
        case [_, "rtf"]:
            with open(file_path) as infile:
                content = infile.read()
                text = rtf_to_text(content)
                print(text)
        case [_, ext]:
            print(f"Ismeretlen formátum: {ext}")

def read_and_print_2(melyik_file):
    file_path = os.path.join("resources", melyik_file)
    file_name, ext = melyik_file.split(".")

    match ext.lower():
        case "txt" | "csv":
            with open(file_path, encoding="utf8") as fp:
                print(fp.read())
        case "docx":
            doc = Document(file_path)
            for para in doc.paragraphs:
                print(para.text)
        case "rtf":
            with open(file_path) as infile:
                content = infile.read()
                text = rtf_to_text(content)
                print(text)
        case other:
            print(f"Ismeretlen formátum: {other}")


def read_and_print_1(melyik_file):
    file_path = os.path.join("resources", melyik_file)
    file_name, ext = melyik_file.split(".")

    if ext == "txt" or ext == "csv":
        with open(file_path, encoding="utf8") as fp:
            print(fp.read())
    elif ext == "docx":
        doc = Document(file_path)
        for para in doc.paragraphs:
            print(para.text)
    elif ext == "rtf":
        with open(file_path) as infile:
            content = infile.read()
            text = rtf_to_text(content)
        print(text)
    else:
        print(f"Ismeretlen formátum: {ext}")

def main():
    print("Nyissunk meg egy fájlt!")
    files = os.listdir("resources")
    for index, a_file in enumerate(files):
        print(f" {index}. {a_file}")
    melyik_file = int(input("Melyik fájlt olvassuk be (sorszám)?"))
    melyik_file = files[melyik_file]
    print(f"Megvan a fájl: {melyik_file}")
    # read_and_print_1(melyik_file)
    # read_and_print_2(melyik_file)
    # read_and_print_3(melyik_file)

if __name__ == '__main__':
    main()
